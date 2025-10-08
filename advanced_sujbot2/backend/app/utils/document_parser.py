"""Utility for parsing document metadata."""
import os
from typing import Dict, Optional
import logging

logger = logging.getLogger(__name__)


def parse_document_metadata(file_path: str) -> Dict[str, int]:
    """
    Parse document and extract metadata.

    Args:
        file_path: Path to document file

    Returns:
        Dict with page_count, word_count, char_count
    """
    file_ext = os.path.splitext(file_path)[1].lower()

    try:
        if file_ext == '.pdf':
            return _parse_pdf(file_path)
        elif file_ext == '.docx':
            return _parse_docx(file_path)
        elif file_ext in ['.txt', '.md']:
            return _parse_text(file_path)
        else:
            logger.warning(f"Unsupported file format: {file_ext}")
            return {"page_count": 1, "word_count": 0, "char_count": 0}
    except Exception as e:
        logger.error(f"Error parsing document {file_path}: {e}")
        return {"page_count": 1, "word_count": 0, "char_count": 0}


def _parse_pdf(file_path: str) -> Dict[str, int]:
    """Parse PDF document."""
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            text = ""

            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + " "

            words = text.split()
            word_count = len(words)
            char_count = len(text)

            return {
                "page_count": page_count,
                "word_count": word_count,
                "char_count": char_count
            }
    except ImportError:
        # Fallback to PyPDF2 if pdfplumber not available
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            page_count = len(reader.pages)
            text = ""

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + " "

            words = text.split()
            word_count = len(words)
            char_count = len(text)

            return {
                "page_count": page_count,
                "word_count": word_count,
                "char_count": char_count
            }
        except Exception as e:
            logger.error(f"Error parsing PDF with PyPDF2: {e}")
            return {"page_count": 1, "word_count": 0, "char_count": 0}


def _parse_docx(file_path: str) -> Dict[str, int]:
    """Parse DOCX document."""
    try:
        from docx import Document

        doc = Document(file_path)

        # Count pages (approximate based on paragraphs)
        page_count = max(1, len(doc.paragraphs) // 30)  # ~30 paragraphs per page

        # Extract text
        text = " ".join([para.text for para in doc.paragraphs])

        words = text.split()
        word_count = len(words)
        char_count = len(text)

        return {
            "page_count": page_count,
            "word_count": word_count,
            "char_count": char_count
        }
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        return {"page_count": 1, "word_count": 0, "char_count": 0}


def _parse_text(file_path: str) -> Dict[str, int]:
    """Parse text document (TXT, MD)."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()

        # Estimate pages (assuming ~500 words per page)
        words = text.split()
        word_count = len(words)
        page_count = max(1, word_count // 500)
        char_count = len(text)

        return {
            "page_count": page_count,
            "word_count": word_count,
            "char_count": char_count
        }
    except Exception as e:
        logger.error(f"Error parsing text file: {e}")
        return {"page_count": 1, "word_count": 0, "char_count": 0}
