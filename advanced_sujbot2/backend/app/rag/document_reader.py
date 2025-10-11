"""
Legal Document Reader - Advanced document parsing with structure awareness.

This module provides comprehensive legal document parsing capabilities including:
- Multi-format support (PDF, DOCX, TXT, XML)
- Structure parsing for Czech laws and contracts
- Reference extraction
- Content classification
"""

import re
import logging
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from abc import ABC, abstractmethod

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from odf import text as odf_text, teletype
    from odf.opendocument import load as odf_load
    HAS_ODF = True
except ImportError:
    HAS_ODF = False

try:
    from striprtf.striprtf import rtf_to_text
    HAS_STRIPRTF = True
except ImportError:
    HAS_STRIPRTF = False

from .models import (
    LegalDocument, DocumentStructure, DocumentMetadata,
    StructuralElement, Part, Chapter, Section, Paragraph,
    Subsection, Letter, Article, Point, LegalReference,
    DocumentType, ElementType
)


# Configure logging
logger = logging.getLogger(__name__)


# ============================================================================
# Custom Exceptions
# ============================================================================

class DocumentReadError(Exception):
    """Base exception for document reading errors"""
    pass


class UnsupportedFormatError(DocumentReadError):
    """Unsupported file format"""
    pass


class StructureParsingError(DocumentReadError):
    """Error parsing document structure"""
    pass


class CorruptedDocumentError(DocumentReadError):
    """Document file is corrupted"""
    pass


# ============================================================================
# Format Readers
# ============================================================================

class FormatReader(ABC):
    """Abstract base class for format-specific readers"""

    @abstractmethod
    async def read(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Read document and return (text, metadata)

        Args:
            file_path: Path to document

        Returns:
            Tuple of (extracted_text, page_info)
        """
        pass


class PDFReader(FormatReader):
    """PDF reader with fallback strategy"""

    async def read(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Read PDF using pdfplumber with PyPDF2 fallback"""

        text = ""
        page_info = {"total_pages": 0, "reader": "none"}

        # Try pdfplumber first (better quality)
        if HAS_PDFPLUMBER:
            try:
                logger.debug(f"Attempting to read PDF with pdfplumber: {file_path}")
                with pdfplumber.open(file_path) as pdf:
                    page_info["total_pages"] = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    page_info["reader"] = "pdfplumber"
                    logger.info(f"Successfully read PDF with pdfplumber: {page_info['total_pages']} pages")
                    return text, page_info
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}, trying PyPDF2 fallback")

        # Fallback to PyPDF2
        if HAS_PYPDF2:
            try:
                logger.debug(f"Attempting to read PDF with PyPDF2: {file_path}")
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    page_info["total_pages"] = len(pdf_reader.pages)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    page_info["reader"] = "pypdf2"
                    logger.info(f"Successfully read PDF with PyPDF2: {page_info['total_pages']} pages")
                    return text, page_info
            except Exception as e:
                logger.error(f"PyPDF2 also failed: {e}")
                raise CorruptedDocumentError(f"Failed to read PDF with all available readers: {e}")

        raise UnsupportedFormatError("No PDF reader available (install pdfplumber or PyPDF2)")


class DOCXReader(FormatReader):
    """DOCX reader using python-docx"""

    async def read(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Read DOCX file"""

        if not HAS_DOCX:
            raise UnsupportedFormatError("python-docx not installed")

        try:
            logger.debug(f"Reading DOCX: {file_path}")
            doc = DocxDocument(file_path)

            # Extract all paragraphs
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

            page_info = {
                "total_pages": 1,  # DOCX doesn't have fixed pages
                "total_paragraphs": len(doc.paragraphs),
                "reader": "python-docx"
            }

            logger.info(f"Successfully read DOCX: {page_info['total_paragraphs']} paragraphs")
            return text, page_info

        except Exception as e:
            logger.error(f"Failed to read DOCX: {e}")
            raise DocumentReadError(f"Error reading DOCX file: {e}")


class ODTReader(FormatReader):
    """ODT (OpenDocument Text) reader using odfpy"""

    async def read(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Read ODT file"""

        if not HAS_ODF:
            raise UnsupportedFormatError("odfpy not installed (pip install odfpy)")

        try:
            logger.debug(f"Reading ODT: {file_path}")
            doc = odf_load(file_path)

            # Extract all text from paragraphs
            paragraphs = doc.getElementsByType(odf_text.P)
            text_parts = []
            for para in paragraphs:
                para_text = teletype.extractText(para)
                if para_text.strip():
                    text_parts.append(para_text)

            text = "\n".join(text_parts)

            page_info = {
                "total_pages": 1,  # ODT doesn't have fixed pages
                "total_paragraphs": len(paragraphs),
                "reader": "odfpy"
            }

            logger.info(f"Successfully read ODT: {page_info['total_paragraphs']} paragraphs")
            return text, page_info

        except Exception as e:
            logger.error(f"Failed to read ODT: {e}")
            raise DocumentReadError(f"Error reading ODT file: {e}")


class RTFReader(FormatReader):
    """RTF reader using striprtf"""

    async def read(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Read RTF file"""

        if not HAS_STRIPRTF:
            raise UnsupportedFormatError("striprtf not installed (pip install striprtf)")

        try:
            logger.debug(f"Reading RTF: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                rtf_content = f.read()

            # Convert RTF to plain text
            text = rtf_to_text(rtf_content)

            page_info = {
                "total_pages": 1,
                "reader": "striprtf"
            }

            logger.info(f"Successfully read RTF: {len(text)} characters")
            return text, page_info

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    rtf_content = f.read()
                text = rtf_to_text(rtf_content)
                logger.info(f"Successfully read RTF with latin-1 encoding")
                page_info = {"total_pages": 1, "reader": "striprtf"}
                return text, page_info
            except Exception as e:
                raise DocumentReadError(f"Error reading RTF file: {e}")

        except Exception as e:
            logger.error(f"Failed to read RTF: {e}")
            raise DocumentReadError(f"Error reading RTF file: {e}")


class TXTReader(FormatReader):
    """Plain text reader"""

    async def read(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Read plain text file"""

        try:
            logger.debug(f"Reading text file: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            page_info = {
                "total_pages": 1,
                "reader": "text"
            }

            logger.info(f"Successfully read text file: {len(text)} characters")
            return text, page_info

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                logger.info(f"Successfully read text file with latin-1 encoding")
                return text, page_info
            except Exception as e:
                raise DocumentReadError(f"Error reading text file: {e}")

        except Exception as e:
            logger.error(f"Failed to read text file: {e}")
            raise DocumentReadError(f"Error reading text file: {e}")


class XMLReader(FormatReader):
    """XML reader"""

    async def read(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Read XML file - basic text extraction"""

        try:
            logger.debug(f"Reading XML file: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            # Basic XML tag removal for text extraction
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text)

            page_info = {
                "total_pages": 1,
                "reader": "xml"
            }

            logger.info(f"Successfully read XML file: {len(text)} characters")
            return text, page_info

        except Exception as e:
            logger.error(f"Failed to read XML: {e}")
            raise DocumentReadError(f"Error reading XML file: {e}")


# ============================================================================
# Structure Parsers
# ============================================================================

class StructureParser(ABC):
    """Abstract base class for structure parsers"""

    @abstractmethod
    async def parse(self, text: str) -> DocumentStructure:
        """Parse document structure from text"""
        pass


class LawStructureParser(StructureParser):
    """Parse Czech law structure"""

    # Regex patterns for Czech laws
    PATTERNS = {
        'law_title': r'Zákon\s+č\.\s*(\d+)/(\d+)\s+Sb\.\s*,?\s*(.*)',
        'part': r'^\s*ČÁST\s+([IVX]+)\s*:?\s*(.*?)$',
        'chapter': r'^\s*HLAVA\s+([IVX]+)\s*:?\s*(.*?)$',
        'section': r'^\s*DÍL\s+([IVX]+)\s*:?\s*(.*?)$',
        'paragraph': r'^\s*§\s*(\d+[a-z]?)\s*(.*)$',
        'subsection': r'^\s*\((\d+)\)\s+(.*)$',
        'letter': r'^\s*([a-z])\)\s+(.*)$',
    }

    async def parse(self, text: str) -> DocumentStructure:
        """Parse law structure"""

        logger.debug("Parsing law structure")
        lines = text.split('\n')

        hierarchy = []
        current_part = None
        current_chapter = None
        current_section = None
        current_paragraph = None
        current_subsection = None

        char_offset = 0

        for line_num, line in enumerate(lines):
            line_length = len(line) + 1  # +1 for newline

            # Try to match structural patterns

            # Part
            if match := re.match(self.PATTERNS['part'], line, re.IGNORECASE):
                number, title = match.groups()
                current_part = Part(
                    element_id=f"part_{number}",
                    element_type=ElementType.PART.value,
                    number=number,
                    title=title.strip(),
                    content="",
                    level=0,
                    parent_id=None,
                    children_ids=[],
                    start_line=line_num,
                    end_line=line_num,
                    start_char=char_offset,
                    end_char=char_offset + line_length,
                    legal_reference=f"Část {number}",
                    metadata={'type': 'part'},
                    chapters=[]
                )
                hierarchy.append(current_part)
                current_chapter = None
                current_paragraph = None
                current_subsection = None
                logger.debug(f"Found part: {current_part.legal_reference}")

            # Chapter
            elif match := re.match(self.PATTERNS['chapter'], line, re.IGNORECASE):
                number, title = match.groups()
                current_chapter = Chapter(
                    element_id=f"chapter_{number}",
                    element_type=ElementType.CHAPTER.value,
                    number=number,
                    title=title.strip(),
                    content="",
                    level=1,
                    parent_id=current_part.element_id if current_part else None,
                    children_ids=[],
                    start_line=line_num,
                    end_line=line_num,
                    start_char=char_offset,
                    end_char=char_offset + line_length,
                    legal_reference=f"Hlava {number}",
                    metadata={'type': 'chapter'},
                    part=current_part,
                    sections=[]
                )
                if current_part:
                    current_part.chapters.append(current_chapter)
                    current_part.children_ids.append(current_chapter.element_id)
                hierarchy.append(current_chapter)
                current_section = None
                current_paragraph = None
                current_subsection = None
                logger.debug(f"Found chapter: {current_chapter.legal_reference}")

            # Section (DÍL)
            elif match := re.match(self.PATTERNS['section'], line, re.IGNORECASE):
                number, title = match.groups()
                current_section = Section(
                    element_id=f"section_{number}",
                    element_type=ElementType.SECTION.value,
                    number=number,
                    title=title.strip(),
                    content="",
                    level=2,
                    parent_id=current_chapter.element_id if current_chapter else None,
                    children_ids=[],
                    start_line=line_num,
                    end_line=line_num,
                    start_char=char_offset,
                    end_char=char_offset + line_length,
                    legal_reference=f"Díl {number}",
                    metadata={'type': 'section'},
                    chapter=current_chapter,
                    paragraphs=[]
                )
                if current_chapter:
                    current_chapter.sections.append(current_section)
                    current_chapter.children_ids.append(current_section.element_id)
                hierarchy.append(current_section)
                current_paragraph = None
                current_subsection = None
                logger.debug(f"Found section: {current_section.legal_reference}")

            # Paragraph
            elif match := re.match(self.PATTERNS['paragraph'], line):
                number, title = match.groups()
                para_num = int(re.sub(r'[a-zA-Z]', '', number))  # Remove letter suffix (both cases)

                # Determine parent (Section if exists, otherwise Chapter)
                parent = current_section if current_section else current_chapter
                parent_id = parent.element_id if parent else None
                level = 3 if current_section else 2

                current_paragraph = Paragraph(
                    element_id=f"paragraph_{number}",
                    element_type=ElementType.PARAGRAPH.value,
                    number=para_num,
                    title=title.strip(),
                    content=title.strip(),
                    level=level,
                    parent_id=parent_id,
                    children_ids=[],
                    start_line=line_num,
                    end_line=line_num,
                    start_char=char_offset,
                    end_char=char_offset + line_length,
                    legal_reference=f"§{number}",
                    metadata={'type': 'paragraph'},
                    chapter=current_chapter,
                    section=current_section,
                    subsections=[],
                    contains_obligation=False,
                    contains_prohibition=False,
                    contains_definition=False
                )

                # Add to parent (Section or Chapter)
                if current_section:
                    current_section.paragraphs.append(current_paragraph)
                    current_section.children_ids.append(current_paragraph.element_id)
                elif current_chapter:
                    # Note: Chapter.sections can contain both Section and Paragraph
                    current_chapter.sections.append(current_paragraph)
                    current_chapter.children_ids.append(current_paragraph.element_id)

                hierarchy.append(current_paragraph)
                current_subsection = None
                logger.debug(f"Found paragraph: {current_paragraph.legal_reference}")

            # Subsection
            elif current_paragraph and (match := re.match(self.PATTERNS['subsection'], line)):
                number, content = match.groups()
                current_subsection = Subsection(
                    element_id=f"subsection_{current_paragraph.number}_{number}",
                    element_type=ElementType.SUBSECTION.value,
                    number=int(number),
                    title=None,
                    content=content.strip(),
                    level=3,
                    parent_id=current_paragraph.element_id,
                    children_ids=[],
                    start_line=line_num,
                    end_line=line_num,
                    start_char=char_offset,
                    end_char=char_offset + line_length,
                    legal_reference=f"§{current_paragraph.number} odst. {number}",
                    metadata={'type': 'subsection'},
                    paragraph=current_paragraph,
                    letters=[]
                )
                current_paragraph.subsections.append(current_subsection)
                current_paragraph.children_ids.append(current_subsection.element_id)
                hierarchy.append(current_subsection)
                logger.debug(f"Found subsection: {current_subsection.legal_reference}")

            # Letter
            elif current_subsection and (match := re.match(self.PATTERNS['letter'], line)):
                letter, content = match.groups()
                letter_elem = Letter(
                    element_id=f"letter_{current_paragraph.number}_{current_subsection.number}_{letter}",
                    element_type=ElementType.LETTER.value,
                    letter=letter,
                    title=None,
                    content=content.strip(),
                    level=4,
                    parent_id=current_subsection.element_id,
                    children_ids=[],
                    start_line=line_num,
                    end_line=line_num,
                    start_char=char_offset,
                    end_char=char_offset + line_length,
                    legal_reference=f"§{current_paragraph.number} odst. {current_subsection.number} písm. {letter}",
                    metadata={'type': 'letter'},
                    subsection=current_subsection
                )
                current_subsection.letters.append(letter_elem)
                current_subsection.children_ids.append(letter_elem.element_id)
                hierarchy.append(letter_elem)
                logger.debug(f"Found letter: {letter_elem.legal_reference}")

            # Regular content - append to current element
            else:
                if current_subsection:
                    current_subsection.content += " " + line.strip()
                    current_subsection.end_line = line_num
                    current_subsection.end_char = char_offset + line_length
                elif current_paragraph:
                    current_paragraph.content += " " + line.strip()
                    current_paragraph.end_line = line_num
                    current_paragraph.end_char = char_offset + line_length

            char_offset += line_length

        # Build structure object
        structure = DocumentStructure(
            hierarchy=hierarchy,
            parts=[e for e in hierarchy if isinstance(e, Part)],
            chapters=[e for e in hierarchy if isinstance(e, Chapter)],
            sections=[e for e in hierarchy if isinstance(e, (Paragraph, Article))],
            all_elements=hierarchy
        )

        logger.info(f"Parsed law structure: {len(structure.parts)} parts, "
                   f"{len(structure.chapters)} chapters, {len(structure.sections)} sections")

        return structure


class ContractStructureParser(StructureParser):
    """Parse contract structure (articles and points)"""

    PATTERNS = {
        'article': r'^\s*(?:Článek|Čl\.|Article)\s+(\d+)\s*:?\s*(.*?)$',
        'point': r'^\s*(\d+)\.(\d+)\s+(.*)$',
        'numbered_item': r'^\s*(\d+)\)\s+(.*)$',
    }

    async def parse(self, text: str) -> DocumentStructure:
        """Parse contract structure"""

        logger.debug("Parsing contract structure")
        lines = text.split('\n')

        hierarchy = []
        current_article = None
        char_offset = 0

        for line_num, line in enumerate(lines):
            line_length = len(line) + 1

            # Article
            if match := re.match(self.PATTERNS['article'], line, re.IGNORECASE):
                number, title = match.groups()
                current_article = Article(
                    element_id=f"article_{number}",
                    element_type=ElementType.ARTICLE.value,
                    number=int(number),
                    title=title.strip(),
                    content=title.strip(),
                    level=0,
                    parent_id=None,
                    children_ids=[],
                    start_line=line_num,
                    end_line=line_num,
                    start_char=char_offset,
                    end_char=char_offset + line_length,
                    legal_reference=f"Článek {number}",
                    metadata={'type': 'article'},
                    points=[]
                )
                hierarchy.append(current_article)
                logger.debug(f"Found article: {current_article.legal_reference}")

            # Point (e.g., 5.1)
            elif current_article and (match := re.match(self.PATTERNS['point'], line)):
                art_num, point_num, content = match.groups()
                point_id = f"{art_num}.{point_num}"
                point_elem = Point(
                    element_id=f"point_{art_num}_{point_num}",
                    element_type=ElementType.POINT.value,
                    number=point_id,
                    title=None,
                    content=content.strip(),
                    level=1,
                    parent_id=current_article.element_id,
                    children_ids=[],
                    start_line=line_num,
                    end_line=line_num,
                    start_char=char_offset,
                    end_char=char_offset + line_length,
                    legal_reference=f"Článek {art_num}.{point_num}",
                    metadata={'type': 'point'},
                    article=current_article
                )
                current_article.points.append(point_elem)
                current_article.children_ids.append(point_elem.element_id)
                hierarchy.append(point_elem)
                logger.debug(f"Found point: {point_elem.legal_reference}")

            # Regular content
            else:
                if current_article:
                    current_article.content += " " + line.strip()
                    current_article.end_line = line_num
                    current_article.end_char = char_offset + line_length

            char_offset += line_length

        # Build structure
        structure = DocumentStructure(
            hierarchy=hierarchy,
            sections=[e for e in hierarchy if isinstance(e, Article)],
            all_elements=hierarchy
        )

        logger.info(f"Parsed contract structure: {len(structure.sections)} articles")

        return structure


class RegulationStructureParser(StructureParser):
    """Parse regulation structure (similar to law but simpler)"""

    async def parse(self, text: str) -> DocumentStructure:
        """Parse regulation - use law parser for now"""
        logger.debug("Parsing regulation structure using law parser")
        law_parser = LawStructureParser()
        return await law_parser.parse(text)


# ============================================================================
# Reference Extractor
# ============================================================================

class ReferenceExtractor:
    """Extract legal references from text"""

    # Reference patterns
    PATTERNS = {
        'paragraph': r'§\s*(\d+[a-z]?)(?:\s+odst\.\s*(\d+))?(?:\s+písm\.\s*([a-z]))?',
        'article': r'[Čč]l(?:ánek|\.)\s*(\d+)(?:\.\s*(\d+))?',
        'law_citation': r'[Zz]ákon(?:a|u)?\s+č\.\s*(\d+)/(\d+)\s*Sb\.',
        'regulation': r'[Vv]yhláška\s+č\.\s*(\d+)/(\d+)\s*Sb\.',
    }

    async def extract(
        self,
        text: str,
        structure: DocumentStructure
    ) -> List[LegalReference]:
        """Extract all references from text"""

        logger.debug("Extracting legal references")
        references = []

        # For each structural element, find references in its content
        for element in structure.all_elements:
            element_refs = self._find_references_in_text(
                element.content,
                element.element_id,
                element.start_char
            )
            references.extend(element_refs)

        logger.info(f"Extracted {len(references)} legal references")
        return references

    def _find_references_in_text(
        self,
        text: str,
        source_element_id: str,
        text_offset: int
    ) -> List[LegalReference]:
        """Find references in a piece of text"""

        references = []

        for ref_type, pattern in self.PATTERNS.items():
            for match in re.finditer(pattern, text):
                ref = self._parse_reference_match(
                    match,
                    ref_type,
                    source_element_id,
                    text_offset,
                    text
                )
                references.append(ref)

        return references

    def _parse_reference_match(
        self,
        match: re.Match,
        ref_type: str,
        source_element_id: str,
        text_offset: int,
        full_text: str
    ) -> LegalReference:
        """Parse a single reference match"""

        # Extract context (50 chars before/after)
        start, end = match.span()
        context_start = max(0, start - 50)
        context_end = min(len(full_text), end + 50)
        context = full_text[context_start:context_end]

        # Parse components based on type
        if ref_type == 'paragraph':
            para, subsec, letter = match.groups()
            target_ref = f"§{para}"
            if subsec:
                target_ref += f" odst. {subsec}"
            if letter:
                target_ref += f" písm. {letter}"

            components = {
                'paragraph': int(re.sub(r'[a-z]', '', para)),
                'subsection': int(subsec) if subsec else None,
                'letter': letter
            }

        elif ref_type == 'article':
            article, point = match.groups()
            target_ref = f"Článek {article}"
            if point:
                target_ref += f".{point}"

            components = {
                'article': int(article),
                'point': point
            }

        elif ref_type == 'law_citation':
            number, year = match.groups()
            target_ref = f"Zákon č. {number}/{year} Sb."
            components = {
                'law_number': number,
                'year': year
            }

        elif ref_type == 'regulation':
            number, year = match.groups()
            target_ref = f"Vyhláška č. {number}/{year} Sb."
            components = {
                'regulation_number': number,
                'year': year
            }

        else:
            target_ref = match.group(0)
            components = {}

        return LegalReference(
            reference_id=f"ref_{source_element_id}_{start}",
            source_element_id=source_element_id,
            source_position=(text_offset + start, text_offset + end),
            target_type=ref_type,
            target_reference=target_ref,
            target_document=None,
            context=context,
            reference_type='direct',
            components=components
        )


# ============================================================================
# Content Classifier
# ============================================================================

class ContentClassifier:
    """Classify content (obligations, prohibitions, definitions)"""

    OBLIGATION_PATTERNS = [
        r'\b(?:musí|je povinen|je třeba|má povinnost)\b',
        r'\b(?:zajistí|provede|předloží)\b',
    ]

    PROHIBITION_PATTERNS = [
        r'\b(?:nesmí|je zakázáno|není dovoleno|není oprávněn)\b',
    ]

    DEFINITION_PATTERNS = [
        r'\b(?:se rozumí|znamená|je definováno jako)\b',
        r'^\s*[A-ZČŘŠŽÝÁÍÉÚŮ][a-zčřšžýáíéúů]+\s+je\b',
    ]

    async def classify(self, structure: DocumentStructure) -> DocumentStructure:
        """Classify content in structural elements"""

        logger.debug("Classifying content")

        for element in structure.all_elements:
            if isinstance(element, Paragraph):
                element.contains_obligation = any(
                    re.search(pattern, element.content, re.IGNORECASE)
                    for pattern in self.OBLIGATION_PATTERNS
                )
                element.contains_prohibition = any(
                    re.search(pattern, element.content, re.IGNORECASE)
                    for pattern in self.PROHIBITION_PATTERNS
                )
                element.contains_definition = any(
                    re.search(pattern, element.content, re.IGNORECASE)
                    for pattern in self.DEFINITION_PATTERNS
                )

        logger.debug("Content classification complete")
        return structure


# ============================================================================
# Main Legal Document Reader
# ============================================================================

class LegalDocumentReader:
    """Main reader orchestrating parsing pipeline"""

    def __init__(self):
        self.format_readers = {
            'pdf': PDFReader(),
            'docx': DOCXReader(),
            'odt': ODTReader(),
            'rtf': RTFReader(),
            'xml': XMLReader(),
            'txt': TXTReader()
        }

        self.structure_parsers = {
            'contract': ContractStructureParser(),
            'law_code': LawStructureParser(),
            'law': LawStructureParser(),  # Alias for law_code
            'regulation': RegulationStructureParser()
        }

        self.reference_extractor = ReferenceExtractor()
        self.content_classifier = ContentClassifier()

    async def read_legal_document(
        self,
        file_path: str,
        document_type: Optional[str] = None
    ) -> LegalDocument:
        """
        Main entry point for reading a legal document

        Args:
            file_path: Path to document
            document_type: 'contract' | 'law_code' | 'regulation' (auto-detect if None)

        Returns:
            LegalDocument with full structure and metadata
        """

        logger.info(f"Reading legal document: {file_path}")

        # 1. Detect format
        file_format = self._detect_format(file_path)
        logger.debug(f"Detected format: {file_format}")

        # 2. Extract raw text
        reader = self.format_readers[file_format]
        raw_text, page_info = await reader.read(file_path)

        # 3. Clean text
        cleaned_text = self._clean_text(raw_text)

        # 4. Auto-detect document type if not provided
        if document_type is None:
            document_type = self._detect_document_type(cleaned_text)
            logger.info(f"Auto-detected document type: {document_type}")

        # 5. Parse structure
        parser = self.structure_parsers[document_type]
        structure = await parser.parse(cleaned_text)

        # 6. Extract references
        references = await self.reference_extractor.extract(
            cleaned_text,
            structure
        )

        # 7. Classify content
        structure = await self.content_classifier.classify(structure)

        # 8. Extract metadata
        metadata = self._extract_metadata(
            file_path,
            cleaned_text,
            structure,
            document_type,
            file_format,
            page_info
        )

        # 9. Create document
        document = LegalDocument(
            document_id=self._generate_document_id(file_path),
            document_type=document_type,
            raw_text=raw_text,
            cleaned_text=cleaned_text,
            structure=structure,
            metadata=metadata,
            references=references,
            parsed_at=datetime.now()
        )

        logger.info(f"Document successfully parsed: {document.document_id}")
        return document

    def _detect_format(self, file_path: str) -> str:
        """Detect file format from extension"""
        ext = Path(file_path).suffix.lower()
        format_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',
            '.odt': 'odt',
            '.rtf': 'rtf',
            '.xml': 'xml',
            '.txt': 'txt',
            '.md': 'txt'
        }
        return format_map.get(ext, 'txt')

    def _detect_document_type(self, text: str) -> str:
        """Auto-detect document type from content"""

        # Law indicators
        law_indicators = [
            r'Zákon\s+č\.\s*\d+/\d+\s+Sb\.',
            r'ČÁST\s+[IVX]+',
            r'HLAVA\s+[IVX]+',
            r'§\s*\d+'
        ]

        # Contract indicators
        contract_indicators = [
            r'Smlouva\s+o\s+\w+',
            r'Článek\s+\d+',
            r'smluvní\s+strany',
            r'dodavatel\s+a\s+objednatel'
        ]

        law_score = sum(
            1 for pattern in law_indicators
            if re.search(pattern, text, re.IGNORECASE)
        )

        contract_score = sum(
            1 for pattern in contract_indicators
            if re.search(pattern, text, re.IGNORECASE)
        )

        if law_score > contract_score:
            return 'law_code'
        elif contract_score > 0:
            return 'contract'
        else:
            return 'regulation'

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""

        # Remove excessive whitespace
        text = re.sub(r' +', ' ', text)

        # Normalize line breaks (but keep paragraph structure)
        text = re.sub(r'\n\n+', '\n\n', text)

        # Remove page numbers (heuristic - standalone numbers on lines)
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)

        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")

        return text.strip()

    def _extract_metadata(
        self,
        file_path: str,
        text: str,
        structure: DocumentStructure,
        document_type: str,
        file_format: str,
        page_info: Dict[str, Any]
    ) -> DocumentMetadata:
        """Extract document metadata"""

        # Get file info
        path_obj = Path(file_path)
        file_size = path_obj.stat().st_size

        # Try to extract title
        title = self._extract_title(text, document_type)

        # Try to extract document number (for laws)
        document_number = None
        if document_type == 'law_code':
            match = re.search(r'Zákon\s+č\.\s*(\d+/\d+)\s+Sb\.', text)
            if match:
                document_number = match.group(1)

        # Count words
        word_count = len(text.split())

        metadata = DocumentMetadata(
            file_path=str(path_obj.absolute()),
            file_format=file_format,
            file_size_bytes=file_size,
            title=title,
            document_number=document_number,
            total_pages=page_info.get('total_pages', 1),
            total_words=word_count,
            total_sections=len(structure.sections)
        )

        return metadata

    def _extract_title(self, text: str, document_type: str) -> str:
        """Extract document title"""

        # Try to find title in first few lines
        lines = text.split('\n')[:10]

        if document_type == 'law_code':
            for line in lines:
                if match := re.search(r'Zákon\s+č\.\s*\d+/\d+\s+Sb\.\s*,?\s*(.*)', line):
                    return match.group(0).strip()

        elif document_type == 'contract':
            for line in lines:
                if 'smlouva' in line.lower():
                    return line.strip()

        # Default: use first non-empty line
        for line in lines:
            if line.strip():
                return line.strip()

        return "Untitled Document"

    def _generate_document_id(self, file_path: str) -> str:
        """Generate unique document ID from file path"""
        return hashlib.md5(file_path.encode()).hexdigest()[:16]
