"""
Document Reader - Modul pro čtení různých formátů dokumentů
Enhanced with batch processing for large-scale document handling
"""

import asyncio
import aiofiles
from pathlib import Path
from typing import Optional, Dict, Any, AsyncGenerator, List, Tuple
import PyPDF2
from docx import Document
import markdown
import logging
from dataclasses import dataclass
import hashlib
from datetime import datetime

logger = logging.getLogger(__name__)


@dataclass
class DocumentBatch:
    """Represents a batch of document content"""
    batch_id: str
    content: str
    page_numbers: List[int]
    metadata: Dict[str, Any]
    byte_offset: int
    byte_length: int


class DocumentReader:
    """Třída pro čtení různých formátů dokumentů s podporou batch processingu"""

    SUPPORTED_FORMATS = {'.pdf', '.docx', '.txt', '.md', '.doc'}
    DEFAULT_BATCH_SIZE = 10  # pages per batch for PDFs
    DEFAULT_CHUNK_SIZE = 50000  # characters per batch for text files

    async def read_document(self, file_path: str) -> str:
        """
        Načte dokument a vrátí jeho textový obsah

        Args:
            file_path: Cesta k souboru

        Returns:
            Textový obsah dokumentu
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Soubor {file_path} neexistuje")

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Nepodporovaný formát souboru: {extension}")

        logger.info(f"Načítání dokumentu: {file_path} (formát: {extension})")

        # Výběr správné metody podle formátu
        if extension == '.pdf':
            return await self._read_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return await self._read_docx(file_path)
        elif extension == '.md':
            return await self._read_markdown(file_path)
        else:  # .txt a ostatní textové formáty
            return await self._read_text(file_path)

    async def _read_pdf(self, file_path: str) -> str:
        """Načte PDF dokument - zkusí více metod pro robustnost"""
        def read_sync():
            text = []

            # Nejdřív zkusíme pdfplumber (robustnější)
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    logger.info(f"PDF má {len(pdf.pages)} stránek (pdfplumber)")
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text.append(f"--- Stránka {i + 1} ---\n{page_text}")

                    if text:
                        return '\n\n'.join(text)
            except Exception as e:
                logger.warning(f"pdfplumber selhal: {e}, zkouším PyPDF2")

            # Pokud pdfplumber selže, zkusíme PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    num_pages = len(pdf_reader.pages)
                    logger.info(f"PDF má {num_pages} stránek (PyPDF2)")

                    for page_num in range(num_pages):
                        try:
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            if page_text:
                                text.append(f"--- Stránka {page_num + 1} ---\n{page_text}")
                        except Exception as e:
                            logger.warning(f"Nelze přečíst stránku {page_num + 1}: {e}")
                            continue

                    if text:
                        return '\n\n'.join(text)
            except Exception as e:
                logger.error(f"PyPDF2 také selhal: {e}")
                raise ValueError(f"Nelze načíst PDF soubor: {e}")

            return '\n\n'.join(text) if text else "Nepodařilo se extrahovat text z PDF"

        # Spuštění synchronní operace v thread poolu
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, read_sync)

    async def _read_docx(self, file_path: str) -> str:
        """Načte DOCX dokument"""
        def read_sync():
            doc = Document(file_path)
            text = []

            # Čtení odstavců
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Detekce nadpisů podle stylu
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name[-1] if paragraph.style.name[-1].isdigit() else '1'
                        text.append(f"\n{'#' * int(level)} {paragraph.text}\n")
                    else:
                        text.append(paragraph.text)

            # Čtení tabulek
            for table in doc.tables:
                text.append("\n[TABULKA]")
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    text.append(row_text)
                text.append("[KONEC TABULKY]\n")

            return '\n'.join(text)

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, read_sync)

    async def _read_markdown(self, file_path: str) -> str:
        """Načte Markdown dokument"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
            # Konverze markdown na plain text, ale zachování struktury
            md = markdown.Markdown()
            html = md.convert(content)
            # Zde bychom mohli použít BeautifulSoup pro lepší extrakci textu
            # Pro jednoduchost vrátíme původní markdown
            return content

    async def _read_text(self, file_path: str) -> str:
        """Načte textový soubor"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            return await f.read()

    async def get_document_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Získá metadata dokumentu

        Args:
            file_path: Cesta k souboru

        Returns:
            Slovník s metadaty
        """
        path = Path(file_path)
        stat = path.stat()

        metadata = {
            'filename': path.name,
            'extension': path.suffix,
            'size_bytes': stat.st_size,
            'size_mb': stat.st_size / (1024 * 1024),
            'created': stat.st_ctime,
            'modified': stat.st_mtime
        }

        # Specifická metadata podle typu
        if path.suffix.lower() == '.pdf':
            try:
                with open(file_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    metadata['pages'] = len(pdf.pages)
                    if pdf.metadata:
                        metadata['pdf_metadata'] = {
                            'title': pdf.metadata.get('/Title', ''),
                            'author': pdf.metadata.get('/Author', ''),
                            'subject': pdf.metadata.get('/Subject', ''),
                            'creator': pdf.metadata.get('/Creator', '')
                        }
            except:
                pass

        return metadata

    async def read_document_batch(
        self,
        file_path: str,
        batch_size: int = None,
        page_range: Tuple[int, int] = None
    ) -> AsyncGenerator[DocumentBatch, None]:
        """
        Čte dokument po dávkách pro efektivní zpracování velkých souborů

        Args:
            file_path: Cesta k souboru
            batch_size: Velikost dávky (počet stránek pro PDF, znaků pro text)
            page_range: Rozsah stránek k načtení (start, end) - pouze pro PDF

        Yields:
            DocumentBatch objekty s částmi dokumentu
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Soubor {file_path} neexistuje")

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Nepodporovaný formát souboru: {extension}")

        logger.info(f"Batch čtení dokumentu: {file_path}")

        # Výběr správné batch metody podle formátu
        if extension == '.pdf':
            async for batch in self._read_pdf_batch(file_path, batch_size or self.DEFAULT_BATCH_SIZE, page_range):
                yield batch
        elif extension in ['.docx', '.doc']:
            async for batch in self._read_docx_batch(file_path, batch_size or self.DEFAULT_CHUNK_SIZE):
                yield batch
        else:  # text files
            async for batch in self._read_text_batch(file_path, batch_size or self.DEFAULT_CHUNK_SIZE):
                yield batch

    async def _read_pdf_batch(
        self,
        file_path: str,
        batch_size: int,
        page_range: Tuple[int, int] = None
    ) -> AsyncGenerator[DocumentBatch, None]:
        """Čte PDF po dávkách stránek"""

        def read_pages_sync(start_page: int, end_page: int):
            text_parts = []
            page_numbers = []

            # Try pdfplumber first
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    total_pages = len(pdf.pages)
                    actual_end = min(end_page, total_pages)

                    for i in range(start_page, actual_end):
                        page = pdf.pages[i]
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"--- Stránka {i + 1} ---\n{page_text}")
                            page_numbers.append(i + 1)

                    return '\n\n'.join(text_parts), page_numbers

            except Exception as e:
                logger.warning(f"pdfplumber batch failed: {e}, trying PyPDF2")

            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    total_pages = len(pdf_reader.pages)
                    actual_end = min(end_page, total_pages)

                    for i in range(start_page, actual_end):
                        try:
                            page = pdf_reader.pages[i]
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(f"--- Stránka {i + 1} ---\n{page_text}")
                                page_numbers.append(i + 1)
                        except Exception as e:
                            logger.warning(f"Cannot read page {i + 1}: {e}")
                            continue

                    return '\n\n'.join(text_parts), page_numbers

            except Exception as e:
                logger.error(f"PyPDF2 batch also failed: {e}")
                raise ValueError(f"Cannot read PDF file: {e}")

        # Get total pages first
        loop = asyncio.get_event_loop()

        def get_page_count():
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    return len(pdf.pages)
            except:
                with open(file_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    return len(pdf.pages)

        total_pages = await loop.run_in_executor(None, get_page_count)

        # Determine page range
        start_page = 0 if not page_range else max(0, page_range[0])
        end_page = total_pages if not page_range else min(total_pages, page_range[1])

        # Process in batches
        batch_num = 0
        for batch_start in range(start_page, end_page, batch_size):
            batch_end = min(batch_start + batch_size, end_page)

            # Read batch
            content, page_nums = await loop.run_in_executor(
                None,
                read_pages_sync,
                batch_start,
                batch_end
            )

            if content:
                batch_id = hashlib.md5(f"{file_path}_{batch_num}".encode()).hexdigest()[:16]

                yield DocumentBatch(
                    batch_id=batch_id,
                    content=content,
                    page_numbers=page_nums,
                    metadata={
                        'file_path': file_path,
                        'batch_number': batch_num,
                        'total_batches': (end_page - start_page + batch_size - 1) // batch_size,
                        'pages_in_batch': len(page_nums),
                        'total_pages': total_pages
                    },
                    byte_offset=batch_start,  # Using page number as offset
                    byte_length=len(content.encode('utf-8'))
                )

            batch_num += 1

    async def _read_docx_batch(
        self,
        file_path: str,
        chunk_size: int
    ) -> AsyncGenerator[DocumentBatch, None]:
        """Čte DOCX po dávkách textu"""

        def read_chunks_sync():
            doc = Document(file_path)
            current_chunk = []
            current_size = 0
            chunks = []

            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_text = paragraph.text
                    para_size = len(para_text)

                    # If adding this paragraph exceeds chunk size, yield current chunk
                    if current_size + para_size > chunk_size and current_chunk:
                        chunks.append('\n'.join(current_chunk))
                        current_chunk = []
                        current_size = 0

                    current_chunk.append(para_text)
                    current_size += para_size

            # Add remaining chunk
            if current_chunk:
                chunks.append('\n'.join(current_chunk))

            return chunks

        loop = asyncio.get_event_loop()
        chunks = await loop.run_in_executor(None, read_chunks_sync)

        byte_offset = 0
        for batch_num, chunk in enumerate(chunks):
            chunk_bytes = chunk.encode('utf-8')
            batch_id = hashlib.md5(f"{file_path}_{batch_num}".encode()).hexdigest()[:16]

            yield DocumentBatch(
                batch_id=batch_id,
                content=chunk,
                page_numbers=[],  # DOCX doesn't have page numbers in the same way
                metadata={
                    'file_path': file_path,
                    'batch_number': batch_num,
                    'total_batches': len(chunks)
                },
                byte_offset=byte_offset,
                byte_length=len(chunk_bytes)
            )

            byte_offset += len(chunk_bytes)

    async def _read_text_batch(
        self,
        file_path: str,
        chunk_size: int
    ) -> AsyncGenerator[DocumentBatch, None]:
        """Čte textový soubor po dávkách"""

        batch_num = 0
        byte_offset = 0

        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            while True:
                chunk = await f.read(chunk_size)
                if not chunk:
                    break

                chunk_bytes = chunk.encode('utf-8')
                batch_id = hashlib.md5(f"{file_path}_{batch_num}".encode()).hexdigest()[:16]

                yield DocumentBatch(
                    batch_id=batch_id,
                    content=chunk,
                    page_numbers=[],
                    metadata={
                        'file_path': file_path,
                        'batch_number': batch_num,
                        'chunk_size': chunk_size
                    },
                    byte_offset=byte_offset,
                    byte_length=len(chunk_bytes)
                )

                byte_offset += len(chunk_bytes)
                batch_num += 1

    async def estimate_processing_time(
        self,
        file_path: str,
        pages_per_second: float = 10.0
    ) -> Dict[str, Any]:
        """
        Odhadne čas zpracování dokumentu

        Args:
            file_path: Cesta k souboru
            pages_per_second: Rychlost zpracování (stránek/sekunda)

        Returns:
            Odhad času a statistiky
        """
        path = Path(file_path)
        stat = path.stat()

        estimate = {
            'file_size_mb': stat.st_size / (1024 * 1024),
            'estimated_time_seconds': 0,
            'estimated_batches': 0
        }

        if path.suffix.lower() == '.pdf':
            try:
                # Get page count
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
            except:
                with open(file_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    page_count = len(pdf.pages)

            estimate['page_count'] = page_count
            estimate['estimated_time_seconds'] = page_count / pages_per_second
            estimate['estimated_batches'] = (page_count + self.DEFAULT_BATCH_SIZE - 1) // self.DEFAULT_BATCH_SIZE

        else:
            # For text files, estimate based on size
            chars_per_second = pages_per_second * 3000  # Assume ~3000 chars per page
            estimate['estimated_characters'] = stat.st_size
            estimate['estimated_time_seconds'] = stat.st_size / chars_per_second
            estimate['estimated_batches'] = (stat.st_size + self.DEFAULT_CHUNK_SIZE - 1) // self.DEFAULT_CHUNK_SIZE

        estimate['estimated_time_minutes'] = estimate['estimated_time_seconds'] / 60

        return estimate

    async def parallel_read(
        self,
        file_paths: List[str],
        max_workers: int = 4
    ) -> AsyncGenerator[Tuple[str, str], None]:
        """
        Paralelně čte více dokumentů

        Args:
            file_paths: Seznam cest k souborům
            max_workers: Maximální počet paralelních čtení

        Yields:
            Tuple (file_path, content) pro každý dokument
        """
        semaphore = asyncio.Semaphore(max_workers)

        async def read_with_semaphore(file_path: str):
            async with semaphore:
                try:
                    content = await self.read_document(file_path)
                    return file_path, content
                except Exception as e:
                    logger.error(f"Failed to read {file_path}: {e}")
                    return file_path, None

        # Create tasks for all files
        tasks = [read_with_semaphore(fp) for fp in file_paths]

        # Process as they complete
        for coro in asyncio.as_completed(tasks):
            result = await coro
            if result[1] is not None:  # Only yield successful reads
                yield result